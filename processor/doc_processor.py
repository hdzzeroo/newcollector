"""
DOC/DOCX 文件处理器
提取 Word 文档的文字内容
"""

import io
from typing import Optional
from dataclasses import dataclass


@dataclass
class DocContent:
    """DOC内容提取结果"""
    success: bool
    text: str
    paragraph_count: int
    error_message: Optional[str] = None


class DocProcessor:
    """DOC/DOCX 文件处理器"""

    def __init__(self, max_paragraphs: int = 50):
        """
        初始化 DOC 处理器

        Args:
            max_paragraphs: 最大提取段落数（模拟"前两页"的效果）
        """
        self.max_paragraphs = max_paragraphs

    def extract_text(self, doc_path: str, max_paragraphs: int = None) -> DocContent:
        """
        提取 DOCX 文件的文字

        Args:
            doc_path: DOCX 文件路径
            max_paragraphs: 最大提取段落数

        Returns:
            DocContent
        """
        max_paragraphs = max_paragraphs or self.max_paragraphs

        # 检查文件扩展名
        if doc_path.lower().endswith('.doc'):
            return self._extract_from_doc(doc_path, max_paragraphs)
        else:
            return self._extract_from_docx(doc_path, max_paragraphs)

    def _extract_from_docx(self, docx_path: str, max_paragraphs: int) -> DocContent:
        """提取 DOCX 文件"""
        try:
            from docx import Document

            doc = Document(docx_path)
            text_parts = []
            count = 0

            for para in doc.paragraphs:
                if count >= max_paragraphs:
                    break
                text = para.text.strip()
                if text:
                    text_parts.append(text)
                    count += 1

            # 也提取表格内容
            for table in doc.tables:
                if count >= max_paragraphs:
                    break
                for row in table.rows:
                    if count >= max_paragraphs:
                        break
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(f"[表格行] {row_text}")
                        count += 1

            full_text = "\n\n".join(text_parts)

            return DocContent(
                success=True,
                text=full_text,
                paragraph_count=len(text_parts)
            )

        except ImportError:
            return DocContent(
                success=False,
                text="",
                paragraph_count=0,
                error_message="需要安装 python-docx: pip install python-docx"
            )
        except Exception as e:
            return DocContent(
                success=False,
                text="",
                paragraph_count=0,
                error_message=f"DOCX 提取失败: {str(e)}"
            )

    def _extract_from_doc(self, doc_path: str, max_paragraphs: int) -> DocContent:
        """
        提取旧版 DOC 文件
        注意：需要额外的库或转换
        """
        try:
            # 尝试使用 antiword (需要系统安装)
            import subprocess
            result = subprocess.run(
                ['antiword', doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                text = result.stdout
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                limited_text = '\n\n'.join(paragraphs[:max_paragraphs])

                return DocContent(
                    success=True,
                    text=limited_text,
                    paragraph_count=min(len(paragraphs), max_paragraphs)
                )
        except FileNotFoundError:
            pass
        except Exception:
            pass

        # 备选方案：提示用户
        return DocContent(
            success=False,
            text="",
            paragraph_count=0,
            error_message="旧版 .doc 格式需要安装 antiword 或转换为 .docx"
        )

    def extract_text_from_bytes(self, doc_bytes: bytes, file_extension: str,
                                 max_paragraphs: int = None) -> DocContent:
        """
        从字节数据提取文字

        Args:
            doc_bytes: 文档字节数据
            file_extension: 文件扩展名 (.doc 或 .docx)
            max_paragraphs: 最大提取段落数

        Returns:
            DocContent
        """
        max_paragraphs = max_paragraphs or self.max_paragraphs

        if file_extension.lower() == '.doc':
            return DocContent(
                success=False,
                text="",
                paragraph_count=0,
                error_message="旧版 .doc 格式不支持从字节直接提取"
            )

        try:
            from docx import Document

            doc = Document(io.BytesIO(doc_bytes))
            text_parts = []
            count = 0

            for para in doc.paragraphs:
                if count >= max_paragraphs:
                    break
                text = para.text.strip()
                if text:
                    text_parts.append(text)
                    count += 1

            full_text = "\n\n".join(text_parts)

            return DocContent(
                success=True,
                text=full_text,
                paragraph_count=len(text_parts)
            )

        except ImportError:
            return DocContent(
                success=False,
                text="",
                paragraph_count=0,
                error_message="需要安装 python-docx: pip install python-docx"
            )
        except Exception as e:
            return DocContent(
                success=False,
                text="",
                paragraph_count=0,
                error_message=f"提取失败: {str(e)}"
            )

    def is_docx_valid(self, docx_path: str) -> bool:
        """检查 DOCX 是否有效"""
        try:
            from docx import Document
            doc = Document(docx_path)
            return True
        except Exception:
            return False


# 测试代码
if __name__ == "__main__":
    processor = DocProcessor(max_paragraphs=30)

    # 测试需要一个实际的 DOCX 文件
    test_path = "test.docx"

    import os
    if os.path.exists(test_path):
        print("测试 DOCX 文本提取...")
        result = processor.extract_text(test_path)

        print(f"\n提取结果:")
        print(f"  成功: {result.success}")
        print(f"  段落数: {result.paragraph_count}")
        print(f"  文本长度: {len(result.text)}")

        if result.text:
            print(f"\n文本预览 (前500字):")
            print(result.text[:500])
    else:
        print(f"测试文件不存在: {test_path}")
        print("请提供一个 DOCX 文件进行测试")
